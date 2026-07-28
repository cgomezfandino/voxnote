"""Tests for the Markdown → Word (.docx) exporter."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from voxnote.pipeline.docx_exporter import markdown_to_docx

SAMPLE = """\
---
tags: [meeting, 2026-06-13]
date: 2026-06-13
time: 10:30
audio: "[[audio/2026-06-13_10-30-00_meeting.wav]]"
---

# 📋 Meeting — meeting

> 🗓️ **Date:** 2026-06-13 · ⏰ **Time:** 10:30 · 🎧 `meeting.wav`

---

## 📝 Summary

We discussed the product launch.

---

## ✅ Decisions

- Launch on July 1
- Hire a designer

---

## 🎯 Action items

- [ ] Prepare the demo @SPEAKER_00 (deadline: 2026-06-20)
- [ ] Review the budget @Ana (deadline: TBD)

---

## 🔜 Next steps

1. Send the minutes
2. Schedule a follow-up

---

<details>
<summary>📄 Full transcript</summary>

[SPEAKER_00]: Hello everyone.

[SPEAKER_01]: Let's begin.

</details>
"""


def _docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_markdown_to_docx_returns_valid_docx() -> None:
    data = markdown_to_docx(SAMPLE)
    assert isinstance(data, bytes)
    assert data[:2] == b"PK"  # a .docx is a ZIP archive
    Document(BytesIO(data))  # re-open to confirm it is structurally valid


def test_markdown_to_docx_includes_key_sections() -> None:
    text = _docx_text(markdown_to_docx(SAMPLE))
    assert "Summary" in text
    assert "We discussed the product launch." in text
    assert "Launch on July 1" in text
    # Emoji are stripped from headings for a clean professional look.
    assert "📝" not in text
    assert "📋" not in text


def test_markdown_to_docx_builds_task_table() -> None:
    document = Document(BytesIO(markdown_to_docx(SAMPLE)))
    assert document.tables, "expected a task table"
    header = [c.text for c in document.tables[0].rows[0].cells]
    assert header == ["Task", "Owner", "Deadline"]
    body_text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    assert "Prepare the demo" in body_text
    assert "SPEAKER_00" in body_text
    assert "2026-06-20" in body_text


def test_markdown_to_docx_sanitizes_invalid_xml_chars() -> None:
    # Control characters (e.g. from a malicious transcript) must not corrupt the docx.
    malicious = "## Summary\n\nText with \x00\x07 control characters."
    text = _docx_text(markdown_to_docx(malicious))
    assert "\x00" not in text and "\x07" not in text
    assert "control characters" in text


def test_markdown_to_docx_flattens_links() -> None:
    # Markdown links become plain label text (the URL, incl. javascript:, is dropped).
    md = "## Summary\n\nSee [the panel](https://example.com/x) and [click](javascript:alert(1)).\n"
    text = _docx_text(markdown_to_docx(md))
    assert "the panel" in text
    assert "click" in text
    assert "https://example.com" not in text
    assert "javascript:" not in text


def test_markdown_to_docx_includes_enriched_sections() -> None:
    # The Word export must carry the speaker-aware enriched sections too.
    md = (
        "# Meeting\n\n"
        "## 👥 Participants\n\n"
        "- **SPEAKER_00** — Led the discussion.\n"
        "- **Ana** — Proposed using JWT.\n\n"
        "## 📌 Key points\n\n- MVP status\n\n"
        "## 💬 Highlights\n\n"
        "> **Ana:** JWT scales better than sessions.\n"
    )
    text = _docx_text(markdown_to_docx(md))
    assert "Participants" in text
    assert "SPEAKER_00" in text
    assert "Ana" in text
    assert "Key points" in text
    assert "JWT scales better than sessions." in text
