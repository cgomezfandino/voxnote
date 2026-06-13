"""Render a Voxnote Markdown note as a Word (.docx) document.

The note Markdown is produced by :mod:`voxnote.pipeline.exporter`, so its structure is
predictable; this converter also handles any historical note (stored only as Markdown).
Headings, lists, task checkboxes and the transcript section become a styled Word document
aimed at non-technical readers.

All user/LLM-derived text is sanitized before insertion (removing characters that are
invalid in XML), which both prevents corrupt ``.docx`` output and is the export-side of
the "sanitize untrusted content before export" hardening.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

_INVALID_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_TASK_RE = re.compile(r"^- \[[ xX]\]\s*(.*)$")
_DEADLINE_RE = re.compile(r"\(deadline:\s*(.*?)\)\s*$")
_OWNER_RE = re.compile(r"@(\S+)\s*$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")


def _clean(text: str) -> str:
    """Strip characters that are invalid in XML (would corrupt the .docx)."""
    return _INVALID_XML.sub("", text)


def _inline(text: str) -> str:
    """Remove Markdown emphasis/code/wikilink syntax for clean Word text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    return text


def _text(value: str) -> str:
    return _clean(_inline(value)).strip()


def _heading_text(value: str) -> str:
    """Strip leading emoji/symbols from a heading for a clean professional look."""
    stripped = re.sub(r"^[\W_]+", "", _text(value)).strip()
    return stripped or _text(value)


def _strip_frontmatter(markdown: str) -> str:
    if markdown.startswith("---"):
        end = markdown.find("\n---", 3)
        if end != -1:
            newline = markdown.find("\n", end + 1)
            return markdown[newline + 1 :] if newline != -1 else ""
    return markdown


def _parse_task(inner: str) -> tuple[str, str, str]:
    deadline = ""
    match = _DEADLINE_RE.search(inner)
    if match:
        deadline = match.group(1).strip()
        inner = inner[: match.start()].rstrip()
    owner = ""
    match = _OWNER_RE.search(inner)
    if match:
        owner = match.group(1).strip()
        inner = inner[: match.start()].rstrip()
    return inner.strip(), owner, deadline


def markdown_to_docx(markdown: str) -> bytes:
    """Convert a Voxnote note (Markdown) to a styled Word document, returned as bytes."""
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    lines = _strip_frontmatter(markdown).split("\n")
    bullets: list[str] = []
    numbered: list[str] = []
    tasks: list[tuple[str, str, str]] = []

    def flush_bullets() -> None:
        for item in bullets:
            document.add_paragraph(_text(item), style="List Bullet")
        bullets.clear()

    def flush_numbered() -> None:
        for item in numbered:
            document.add_paragraph(_text(item), style="List Number")
        numbered.clear()

    def flush_tasks() -> None:
        if not tasks:
            return
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ("Tarea", "Responsable", "Fecha límite")):
            cell.text = label
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for tarea, owner, deadline in tasks:
            cells = table.add_row().cells
            cells[0].text = _text(tarea)
            cells[1].text = _text(owner) or "—"
            cells[2].text = _text(deadline) or "—"
        tasks.clear()
        document.add_paragraph()

    def flush_all() -> None:
        flush_bullets()
        flush_numbered()
        flush_tasks()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # Transcript collapsible block → heading + paragraphs.
        if stripped.startswith("<details"):
            flush_all()
            title = "Transcripción completa"
            i += 1
            inner: list[str] = []
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith("<summary>"):
                    title = _heading_text(re.sub(r"</?summary>", "", line))
                    i += 1
                    continue
                if line == "</details>":
                    i += 1
                    break
                inner.append(lines[i])
                i += 1
            document.add_heading(title or "Transcripción completa", level=1)
            for block in "\n".join(inner).strip().split("\n"):
                if block.strip():
                    document.add_paragraph(_text(block))
            continue

        if stripped.startswith("</") or stripped in ("<summary>", "</summary>"):
            i += 1
            continue

        task_match = _TASK_RE.match(stripped)
        if task_match:
            flush_bullets()
            flush_numbered()
            tasks.append(_parse_task(task_match.group(1)))
            i += 1
            continue

        if stripped.startswith("### "):
            flush_all()
            document.add_heading(_heading_text(stripped[4:]), level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_all()
            document.add_heading(_heading_text(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith("# "):
            flush_all()
            document.add_heading(_heading_text(stripped[2:]), level=0)
            i += 1
            continue

        bullet_match = _BULLET_RE.match(stripped)
        if bullet_match:
            flush_numbered()
            flush_tasks()
            bullets.append(bullet_match.group(1))
            i += 1
            continue

        numbered_match = _NUMBERED_RE.match(stripped)
        if numbered_match:
            flush_bullets()
            flush_tasks()
            numbered.append(numbered_match.group(1))
            i += 1
            continue

        if stripped.startswith("> "):
            flush_all()
            paragraph = document.add_paragraph()
            paragraph.add_run(_text(stripped[2:])).italic = True
            i += 1
            continue

        # Blank lines and horizontal rules (---, ***, ___) act as separators.
        if set(stripped) <= {"-", "*", "_"}:
            flush_all()
            i += 1
            continue

        flush_all()
        document.add_paragraph(_text(stripped))
        i += 1

    flush_all()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
